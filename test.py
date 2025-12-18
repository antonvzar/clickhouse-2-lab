import time
from statistics import mean
import datetime
from clickhouse_driver import Client
from docx import Document

LOG_TXT = None
DOC = None

def log(message: str = "") -> None:
    print(message)
    if LOG_TXT is not None:
        LOG_TXT.write(message + "\n")
    if DOC is not None:
        DOC.add_paragraph(message)

CLICKHOUSE_HOST = "localhost"
CLICKHOUSE_PORT = 9000       
CLICKHOUSE_DB = "ecom"       
ITERATIONS = 30              

client = Client(
    host="localhost",
    port=9000,
    database="ecom",
    user="benchmark",
    password=""
)



# Querys

# Top-20 by number of products in categories
RAW_TOP_CATEGORIES = """
SELECT
    category_id,
    count() AS offers_cnt
FROM ecom_offers
GROUP BY category_id
ORDER BY offers_cnt DESC
LIMIT 20
"""

MV_TOP_CATEGORIES = """
SELECT
    category_id,
    offers_cnt
FROM catalog_by_category_mv
ORDER BY offers_cnt DESC
LIMIT 20
"""


# Top-30 by number of products in brands
RAW_TOP_BRANDS = """
SELECT
    vendor,
    count() AS offers_cnt
FROM ecom_offers
GROUP BY vendor
ORDER BY offers_cnt DESC
LIMIT 30
"""

MV_TOP_BRANDS = """
SELECT
    vendor,
    sum(offers_cnt) AS offers_cnt
FROM catalog_by_brand_mv
GROUP BY vendor
ORDER BY offers_cnt DESC
LIMIT 30
"""


# Average number of products per brand in each category
RAW_AVG_OFFERS_PER_BRAND = """
SELECT
    category_id,
    avg(offers_per_brand) AS avg_offers_per_brand
FROM
(
    SELECT
        category_id,
        vendor,
        count() AS offers_per_brand
    FROM ecom_offers
    GROUP BY category_id, vendor
)
GROUP BY category_id
ORDER BY avg_offers_per_brand DESC
"""

MV_AVG_OFFERS_PER_BRAND = """
SELECT
    category_id,
    avg(offers_cnt) AS avg_offers_per_brand
FROM catalog_by_brand_mv
GROUP BY category_id
ORDER BY avg_offers_per_brand DESC
"""


# Analysis of products without events through raw_events
RAW_OFFERS_WITHOUT_EVENTS = """
SELECT
    o.offer_id,
    o.category_id,
    o.vendor,
    o.price
FROM ecom_offers AS o
LEFT JOIN
(
    SELECT DISTINCT ContentUnitID AS offer_id
    FROM raw_events
) AS e
    ON o.offer_id = e.offer_id
WHERE e.offer_id IS NULL
"""

# Analysis of products without events through MV
MV_OFFERS_WITHOUT_EVENTS = """
SELECT
    o.offer_id,
    o.category_id,
    o.vendor,
    o.price
FROM ecom_offers AS o
LEFT JOIN offer_events_mv AS ev
    ON o.offer_id = ev.offer_id
WHERE ev.offer_id IS NULL
"""


QUERIES = {
    "raw_top_categories": RAW_TOP_CATEGORIES,
    "mv_top_categories": MV_TOP_CATEGORIES,

    "raw_top_brands": RAW_TOP_BRANDS,
    "mv_top_brands": MV_TOP_BRANDS,

    "raw_avg_offers_per_brand": RAW_AVG_OFFERS_PER_BRAND,
    "mv_avg_offers_per_brand": MV_AVG_OFFERS_PER_BRAND,

    "raw_offers_without_events": RAW_OFFERS_WITHOUT_EVENTS,
    "mv_offers_without_events": MV_OFFERS_WITHOUT_EVENTS,
}


def run_benchmark(name: str, sql: str, iterations: int) -> None:
    times = []

    log(f"\n=== Query {name} ===")
    client.execute(sql)

    for i in range(iterations):
        t0 = time.perf_counter()
        client.execute(sql)
        dt = time.perf_counter() - t0
        times.append(dt)
        log(f"  iteration {i + 1:2d}/{iterations}: {dt:.4f} s")

    log(f"\nResults for {name}:")
    log(f"  min   = {min(times):.4f} s")
    log(f"  mean  = {mean(times):.4f} s")
    log(f"  max   = {max(times):.4f} s")
    log("-" * 40)



def main():
    global LOG_TXT, DOC

    # Имя файлов со штампом времени, чтобы не перезатирать результаты
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = f"clickhouse_load_test_{timestamp}"

    # Открываем txt-лог
    LOG_TXT = open(base_name + ".txt", "w", encoding="utf-8")

    # Готовим Word-документ, если библиотека доступна
    if Document is not None:
        DOC = Document()
        DOC.add_heading("ClickHouse load test", level=1)
    else:
        DOC = None

    log("Starting ClickHouse load test\n")
    log(f"Host: {CLICKHOUSE_HOST}:{CLICKHOUSE_PORT}, database: {CLICKHOUSE_DB}")
    log(f"Number of iterations per query: {ITERATIONS}\n")

    # Гоним все запросы
    for name, sql in QUERIES.items():
        run_benchmark(name, sql, ITERATIONS)

    # Закрываем txt
    LOG_TXT.close()

    # Сохраняем Word, если делали
    if DOC is not None:
        DOC.save(base_name + ".docx")

    log(f"\nResults saved to: {base_name}.txt" + (" and .docx" if DOC is not None else ""))



if __name__ == "__main__":
    main()
